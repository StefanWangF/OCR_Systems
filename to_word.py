# coding=utf-8
import os
import json
from docx.shared import RGBColor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.text import WD_COLOR_INDEX
import re
from pypinyin import pinyin, Style


class ToWord:
    def __init__(self):
        # 文件绝对路径
        self.absolute_path = os.path.dirname(os.getcwd())
        self.generate_file_dir = self.absolute_path + r"\成品"
        self.path_dir = self.absolute_path + r"\录入目录.txt"
        if not os.path.exists(self.generate_file_dir):  # 检查目录是否存在
            os.makedirs(self.generate_file_dir)  # 创建目录
        self.path_name = self.read_path(filename=self.path_dir)
        self.generate_file_dir_word = self.absolute_path + r"\成品\word"
        if not os.path.exists(self.generate_file_dir_word):  # 检查目录是否存在,创建word存储目录
            os.makedirs(self.generate_file_dir_word)  # 创建目录
        self.word_dir = self.generate_file_dir_word + rf"\{self.path_name['pfd_filename']}.docx"

    def make_word_str(self, txt_page):
        """
        传入参数为字典，将字典处理为需要写入word的内容
        :param txt_page: {"txt文件名字": "txt内容", "txt文件名字": "txt内容", ...}
        :return: {"正文": {"中文目录": str, "英文目录": str, ...}, "注释": {"注释": str}}
        """
        global english_title_list
        english_title_list = []
        all_dict = {"正文": {}, "注释": {}}
        make_txt_zhushi_list = []
        new_txt_page = {}
        new_txt_page["中文目录"] = "\n".join(txt_page["中文目录"].split("\n")[1:])
        lines = new_txt_page["中文目录"].split('\n')
        non_empty_lines = [line for line in lines if len(line.strip()) > 0]
        all_dict["正文"]["中文目录"] = '\n'.join(non_empty_lines)
        del txt_page["中文目录"]

        new_txt_page["英文目录"] = "\n".join(txt_page["英文目录"].split("\n")[1:])
        lines = new_txt_page["英文目录"].split('\n')
        non_empty_lines = [line for line in lines if len(line.strip()) > 0]
        all_dict["正文"]["英文目录"] = '\n'.join(non_empty_lines)
        del txt_page["英文目录"]
        # print(txt_page["党的十八大以来党领导全面依法治国的重大创新"])
        # print(self.use_re(txt_str=txt_page["党的十八大以来党领导全面依法治国的重大创新"], re_expression=r'作者＝(.*?)作者单位＝'))
        for key, value in txt_page.items():
            temp_dict = {}
            # 取作者名
            temp_dict["作者"] = "　　" + self.use_re(txt_str=value, re_expression=r'作者＝(.*?)作者单位＝')
            # 取作者单位
            temp_dict["作者单位"] = "　　" + self.use_re(txt_str=value, re_expression=r'作者单位＝(.*?)摘要＝')
            # 取基金项目
            temp_dict["基金项目"] = "　　" + self.use_re(txt_str=value, re_expression=r'基金项目＝(.*?)中图分类号＝')
            # 取中文摘要
            temp_dict["中文摘要"] = "　　" + "摘要：" + str(self.use_re(txt_str=value, re_expression=r'摘要＝(.*?)英文摘要＝'))
            # 取关键字
            temp_dict["关键字"] = "　　" + "关键词：" + str(self.use_re(txt_str=value, re_expression=r'关键字＝(.*?)英文关键字＝'))
            # 取中图分类号
            temp_dict["中图分类号"] = "　　" + "中图分类号：" + self.use_re(txt_str=value, re_expression=r'中图分类号＝(.*?)文献标识码＝')
            # 取文献标识码
            temp_dict["文献标识码"] = "　　" + "文献标识码：" + self.use_re(txt_str=value, re_expression=r'文献标识码＝(.*?)文章编号＝')
            # 取文章编号
            temp_dict["文章编号"] = "　　" + "文章编号：" + self.use_re(txt_str=value, re_expression=r'文章编号＝(.*?)内容＝')
            # 取内容
            temp_dict["内容"] = "　　" + self.use_re(txt_str=value, re_expression=r'内容＝(.*?)注释＝')
            # 取英文标题
            temp_dict["英文标题"] = "　　" + self.use_re(txt_str=value, re_expression=r'英文标题＝(.*?)副标题＝')
            english_title_list.append(temp_dict["英文标题"])
            # 从英文目录中提取英文作者名
            english_dir = new_txt_page["英文目录"]
            # print(english_dir)
            english_dir_list = english_dir.split("\n")
            english_name = ""
            # print(temp_dict)
            # for i in english_dir_list:
            #     print(i)
            #     if str(temp_dict["英文标题"]).lstrip() in i:
            #         name_line = english_dir_list[int(int(english_dir_list.index(i)) + 1)]
            # result1 = re.search(r'^(.*?)\s\d*$', name_line)
            # if result1:
            #     extracted_text1 = result1.group(1)
            # else:
            #     extracted_text1 = ""
            # english_name = extracted_text1.replace("…", "")
            pinyin_list = pinyin(temp_dict["作者"], style=Style.NORMAL)
            # 提取姓和名的拼音，并在姓后面加一个空格
            last_name = pinyin_list[0][0]
            first_name = ''.join([p[0] for p in pinyin_list[1:]])
            # 构建转换后的全拼名
            converted_name = last_name + ' ' + first_name
            if "," in english_name:
                english_name =english_name.replace(",", "\n   ", 1)
            temp_dict["英文作者名"] = converted_name
            # 取英文摘要
            temp_dict["英文摘要"] = "　　" + "Abstract: " + str(self.use_re(txt_str=value, re_expression=r'英文摘要＝(.*?)关键字＝'))
            # 取英文关键字
            temp_dict["英文关键字"] = "　　" + "Keywords: " + str(self.use_re(txt_str=value, re_expression=r'英文关键字＝(.*?)基金项目＝'))
            # 末尾外加基金项目,即注释第一行
            temp_dict["end"] = temp_dict["基金项目"]
            # print(temp_dict)
            # 整合标题正文
            new_txt_page.update(temp_dict)
            # 处理注释,将所有的注释放在一个列表中
            txt_zhushi = self.use_re(txt_str=value, re_expression=r'注释＝(.*?)参考文献＝')
            # 判断注释第一行是否是[1]，如果不是则剔除
            if re.search(r'\[\d+\]', txt_zhushi.split("\n")[0]):
                txt_zhushi_list = txt_zhushi.split("\n")
            else:
                txt_zhushi_list = txt_zhushi.split("\n")[1:]
            # 删除空行
            for i in txt_zhushi_list:
                result = re.search(r'\[\d+\]', i)
                if result:
                    make_txt_zhushi_list.append(i.lstrip())
            all_dict["正文"][key] = "\n".join(list(temp_dict.values()))

        # 遍历每个段落并进行替换
        counter = 1
        for i in range(len(make_txt_zhushi_list)):
            make_txt_zhushi_list[i] = re.sub(r'\[\d+\]', lambda match: f'[{counter}]', make_txt_zhushi_list[i], count=1)
            counter += 1
        # 整理好的注释，给注释添加缩进
        make_txt_zhushi_str = "\n".join("　　" + paragraph for paragraph in make_txt_zhushi_list)
        all_dict["注释"] = {"注释": make_txt_zhushi_str}
        # print(all_dict)
        return all_dict

    def make_word_file(self, message_dict):
        """
        制作word成品，将传入的数据，分割为正文和注释部分
        将所有[数字]转换为上标的形式
        先写入中英文目录，再按照目录章节顺序，给标题加标题1格式，加粗，并调整行间距，给正文调整首行缩进
        写入word中
        :param message_dict: {"正文": {"中文目录": {}, "英文目录": {}, ...}, "注释": {"注释": []}}
        :return:无返回值，直接生成文件
        """
        # 创建Document对象
        bold_strings = ["Keywords:", "Abstract:"] + english_title_list

        doc = Document()

        for k, v in message_dict.items():
            for v_k, v_v in v.items():
                word_title = v_k
                word_para = v_v
                if word_title != "注释":
                    title = doc.add_heading(level=1)
                    run = title.add_run(word_title)
                    run.bold = True
                    run.font.size = Pt(22)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                    title_paragraph_format = title.paragraph_format
                    title_paragraph_format.line_spacing = 2.0
                    title_paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    title_paragraph_format.space_before = Pt(6)

                para = doc.add_paragraph()
                last_match_end = 0
                matches = list(re.finditer(r'\[\d+\]', word_para))

                for match in matches:
                    run = para.add_run(word_para[last_match_end:match.start()])
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                    run = para.add_run(word_para[match.start():match.end()])
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.superscript = True

                    last_match_end = match.end()

                run = para.add_run(word_para[last_match_end:])
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                para_format = para.paragraph_format
                para_format.space_before = Pt(0)
                para_format.space_after = Pt(0)
                para_format.line_spacing = 1.0
                para_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        # 保存文档
        doc.save(self.word_dir)

    def read_txt_page(self, path_name):
        """
        标题
        作者
        作者单位
        基金项目
        摘要
        关键词
        中图分类号
        文献标识码
        文章编号
        内容
        英文标题
        作者名英文
        英文关摘要
        英文关键字
        基金项目
        此函数用来将传入文件名，拼接为绝对路径，去读取txt内容，返回字典
        :param path_name: 此变量为需要读取的txt文件名，不包括路径和后缀，只是名字
        :return:返回字典{"txt文件名字": "txt内容", "txt文件名字": "txt内容", ...}
        """
        dict_path_page = {}     # 获取到对于标题的txt内容，格式为{"标题":"txt内容",....}
        for i in path_name["keys"]:
            txt_path = self.generate_file_dir + "\\" + i + r".txt"
            dict_path_page[i] = self.read_txt(filename=txt_path)
        return dict_path_page

    def read_path(self, filename):
        """
        读取目录txt，返回dict
        :param filename: 传入文件名
        :return: 返回值为目录txt中的key值与源文件名，即：中文目录，英文目录，标题1，标题2...
        {'keys': dict_keys(['中文目录', '英文目录',。。。]), 'pfd_filename': '上海政法学院学报2023年第1期'}
        """
        with open(filename, 'r', encoding='utf-8') as f:
            path = eval(f.read())
            yuan_file_dir = path["源文件绝对路径"]
            # 使用os.path.basename获取文件名（包括扩展名）
            filename_with_extension = os.path.basename(yuan_file_dir)
            # 使用os.path.splitext分离文件名和扩展名
            pdf_filename, extension = os.path.splitext(filename_with_extension)
            del path["源文件绝对路径"]
            return {"keys": path.keys(), "pfd_filename": pdf_filename}

    def read_txt(self, filename):
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read()

    def use_re(self, txt_str, re_expression):
        """
        传入需要匹配的str和正则表达式
        :param txt_str:
        :param re_expression:
        :return: 匹配到的字符
        """
        pattern = re.compile(re_expression, re.S)  # re.S让.匹配包括换行在内的所有字符
        # 使用正则表达式提取内容
        match = re.search(pattern, txt_str)
        if match:
            author = match.group(1).strip()  # 使用group(1)获取第一个括号的内容，并使用strip()去掉可能存在的前后空白符
        else:
            author = ""
        return author


if __name__ == '__main__':
    to_word = ToWord()
    # print(to_word.path_name)
    txt_page = to_word.read_txt_page(path_name=to_word.path_name)
    make_word_str = to_word.make_word_str(txt_page=txt_page)
    to_word.make_word_file(message_dict=make_word_str)
